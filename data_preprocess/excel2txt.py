import pandas as pd
from pyquery import PyQuery as pq
from docx import Document
import os
from tqdm import tqdm


output_dir_docx = 'output_docx_项目招投标_一带一路'
os.makedirs(output_dir_docx, exist_ok=True)
output_dir_txt = 'output_txt_项目招投标_一带一路'
os.makedirs(output_dir_txt, exist_ok=True)

write_fail = open('write_fail-中标.txt', 'w', encoding='utf-8')
write_fail.write('{}'.format('title') + '\n')

df = pd.read_excel('项目招投标-一带一路.xlsx').astype(str)
for i in tqdm(range(len(df))):
    # if '中标' in df['title'][i]:
    try:
        content = df['content'][i]
        content = pq(content).text()
        # 写入.docx文件
        doc = Document()
        doc.add_heading(df['title'][i], level=1)
        doc.add_paragraph(content)
        docx_file_name = '{}.docx'.format(df['title'][i])
        doc.save(os.path.join(output_dir_docx, docx_file_name))

        # 写入.txt文件
        txt_file_name = '{}.txt'.format(df['title'][i])
        with open(os.path.join(output_dir_txt, txt_file_name), 'w', encoding='utf-8') as f:
            f.write(df['title'][i] + '\n')
            f.write(content + '\n\n')
    except Exception as e:
        print('Error: {}, title={} 写入失败！skip +1'.format(e, df['title'][i]))
        write_fail.write('{}'.format(df['title']) + '\n')
