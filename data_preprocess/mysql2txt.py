import pymysql
import pandas as pd
from docx import Document
import os
from tqdm import tqdm
from pyquery import PyQuery as pq
import re


def split_text(html):
    """
    把文章切分成段落
    :param html:
    :return:
    """
    try:
        content_clean = pq(html).text()
        content_clean_spaceline = re.sub('\n+', '\n', content_clean.replace('\t', '').replace('\r', ''))  # 去掉多余的空行
        paragraphs_split = [p for p in content_clean_spaceline.split('\n') if len(p) >= 0]
    except Exception as e:
        print('Error: %s, skip +1' % e)
        paragraphs_split = []
    return paragraphs_split


output_dir_docx = 'output_docx'
os.makedirs(output_dir_docx, exist_ok=True)
output_dir_txt = 'output_txt'
os.makedirs(output_dir_txt, exist_ok=True)

# 连接MySQL数据库
mydb = pymysql.connect(
    host="114.116.11.225",
    port=3306,
    user="root",
    password="@zzsn9988",
    database="auditNew"
)

print('query ...')
# 读取sys_core_law表的数据
df = pd.read_sql("SELECT id, title, pro_content as content FROM sys_core_law "
                 "WHERE pro_content is not null limit 20000", con=mydb)

print('write ...')
# 将数据写入.docx文件和.txt文件
write_fail = open('write_fail.txt', 'w', encoding='utf-8')
write_fail.write('{},{}'.format('id', 'title') + '\n')
for i in tqdm(range(len(df))):
    try:
        content = df['content'][i]
        content = pq(content).text()
        # 写入.docx文件
        doc = Document()
        doc.add_heading(df['title'][i], level=1)
        doc.add_paragraph(content)
        docx_file_name = '{}.docx'.format(df['title'][i])
        doc.save(os.path.join(output_dir_docx, docx_file_name))

        # 写入.txt文件
        txt_file_name = '{}.txt'.format(df['title'][i])
        with open(os.path.join(output_dir_txt, txt_file_name), 'w', encoding='utf-8') as f:
            f.write(df['title'][i] + '\n')
            f.write(content + '\n\n')
    except Exception as e:
        print('Error: {}, id={}, title={} 写入失败！skip +1'.format(e, df['id'][i], df['title'][i]))
        write_fail.write('{},{}'.format(df['id'][i], df['title']) + '\n')

print('Finished !')
