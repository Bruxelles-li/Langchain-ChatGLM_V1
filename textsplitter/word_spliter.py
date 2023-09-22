from docx import Document
import numpy as np
import re


class TextCut:
    def __init__(self, min_len=20, step=10, stop_list=None):
        self.min_len = min_len  # 自定义最短长度
        self.step = step  # 自定义划窗步长
        if stop_list and isinstance(stop_list, list):
            self.stop_list = stop_list  # 自定义分割标点符
        else:
            self.stop_list = ['.', '!', '|', '。', '！', '；', ';', '?', '？', ',']
        self.split_patten = '[' + ''.join(self.stop_list) + ']'

    def find_now_index(self, now_point, sum_len_list):
        for i in range(len(sum_len_list) - 1):
            if now_point >= sum_len_list[i] and now_point < sum_len_list[i + 1]:
                return i + 1
        else:
            return 0

    def cut(self, text):
        if not isinstance(text, str):
            raise TypeError
        split_text = re.split(self.split_patten, text)
        len_list = np.array([len(x) for x in split_text])
        sum_len_list = np.cumsum(len_list)
        result_list = []
        end_point = 0
        pre_index = 0
        while end_point <= sum_len_list[-1]:
            end_point += self.step
            now_index = self.find_now_index(end_point, sum_len_list)
            if np.sum(len_list[pre_index:now_index]) >= self.min_len:
                result_list.append(''.join(split_text[pre_index:now_index]))
                pre_index = now_index

        if pre_index < len(split_text):
            last = ''.join(split_text[pre_index:])
            if len(last) >= self.min_len:
                result_list.append(last)
        return result_list


def split_paragraph(document):
    text_cut_helper = TextCut(min_len=100, step=10, stop_list=['。'])
    paragraphs = document.paragraphs
    for paragraph in paragraphs:
        # 获取段落的文本内容
        text = paragraph.text
        # 将文本内容按照指定的分隔符分割成若干段落
        if len(text) > 500:
            text_list = text_cut_helper.cut(text)
            for i in text_list:
                print(i)


# 创建一个空白文档
document = Document('/home/zzsn/zhangtao/pycharm_projects/langchain/langchain-ChatGLM/knowledge_base/国资国企研究知识库-test/content/wKjIbGSmpr2ABmfyABXCmmR9Zyo746.doc')

# # 添加一个段落
# document.add_paragraph('这是一个示例段落，可以根据需要进行分割。')

# 将第一个段落分成若干段落
split_paragraph(document)

# 保存文档
document.save('output.docx')
